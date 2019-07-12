# coding=utf-8

import os
import sys
from typing import Dict, Any

import docx
import re
from pymongo import MongoClient


def process1_1(process):
    only_check = True if process['only_check'].upper() == 'YES' else False
    if not only_check:
        client = MongoClient(host='127.0.0.1', port=27017)
        db = client['emr']
        col = db['zju4h']
        if process['drop_collections_before_save'].upper() == 'YES':
            col.delete_many({})

    for dir in process['source_dir']:
        print("Reading all files from {}".format(dir))
        files = os.listdir(dir)
        #files = ['138060.docx']
        count = 0
        for file in files:
            if not file.endswith('.docx'):
                continue
            count = count + 1
            print("[{}] Reading {}".format(count, file))
            doc = read_docx(os.path.join(dir, file))
            col.insert_one(doc)


def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = ''
    for para in doc.paragraphs:
        full_text = full_text + para.text + '\n'

    # 将full_text分成三部份admission_record_text、progress_note_text、discharge_record_text

    regex = '入  院  记  录.*'
    r = re.search(regex, full_text, re.S)
    admission_record_text = r.group() if r is not None else ''
    admission_record_text = admission_record_text.split('首次病程记录')[0]
    #admission_record_text = admission_record_text.split('出  院  记  录')[0]
    admission_record_text = re.split('(出  院  记  录)|(入院情况.*住院经过)', admission_record_text, 1, re.S)[0]

    regex = '首次病程记录.*'
    r = re.search(regex, full_text, re.S)
    progress_note_text = r.group() if r is not None else ''
    progress_note_text = progress_note_text.split('入  院  记  录')[0]
    #progress_note_text = progress_note_text.split('出  院  记  录')[0]
    progress_note_text = re.split('(出  院  记  录)|(入院情况.*住院经过)', progress_note_text, 1, re.S)[0]


    regex = '死亡记录.*'
    r = re.search(regex, full_text, re.S)
    if r is not None:
        death_record_text = r.group()
        death_record_text = death_record_text.split('首次病程记录')[0]
        death_record_text = death_record_text.split('入  院  记  录')[0]
        alive = False
    else:
        regex = '(出  院  记  录.*)|(入院情况.*住院经过.*)'
        r = re.search(regex, full_text, re.S)
        discharge_record_text = r.group() if r is not None else ''
        discharge_record_text = discharge_record_text.split('首次病程记录')[0]
        discharge_record_text = discharge_record_text.split('入  院  记  录')[0]
        alive = True

    # 判断入院记录和出院记录的先后顺序

    if alive:
        regex = '入  院  记  录.*((出  院  记  录.*)|(入院情况.*住院经过))'
        r = re.search(regex, full_text, re.S)
        if r is not None:
            t1 = 0
            t2 = 1
        else:
            t1 = 1
            t2 = 0
    else:
        t1 = 0

    admission_record = dict()
    progress_note = dict()
    if alive:
        discharge_record = dict()
    else:
        death_record = dict()

    table = doc.tables[t1]
    admission_record['姓名'] = table.cell(0, 1).text
    admission_record['职业'] = table.cell(0, 3).text
    admission_record['性别'] = table.cell(1, 1).text
    admission_record['工作单位'] = table.cell(1, 3).text
    admission_record['出生日期'] = table.cell(2, 1).text
    admission_record['户口地址'] = table.cell(2, 3).text
    admission_record['婚姻'] = table.cell(3, 1).text
    admission_record['联系电话'] = table.cell(3, 3).text
    admission_record['出生地'] = table.cell(4, 1).text
    admission_record['入院时间'] = table.cell(4, 3).text
    admission_record['民族'] = table.cell(5, 1).text
    admission_record['病史陈述者'] = table.cell(5, 3).text

    if alive:
        table = doc.tables[t2]
        if len(table.rows) == 3:
            i = 0
            c = 0
        elif len(table.rows) == 4:
            i = 1
            c = 0
        else:
            # 如果有5行，出院日期和出院诊断是在第5列，用c变量来调整
            i = 2
            c = 1

        discharge_record['入院日期'] = table.cell(i, 1).text
        discharge_record['出院日期'] = table.cell(i, c + 3).text
        discharge_record['入院诊断'] = table.cell(i + 1, 1).text
        discharge_record['出院诊断'] = table.cell(i + 1, c + 3).text
        discharge_record['住院天数'] = table.cell(i + 2, 1).text

    # 处理入院记录

    regex = '主诉\(Chief complaint\)：(?P<主诉>.+)\n' \
            '现病史（History of Present Illness）:(?P<现病史>.+)\n' \
            '既往史（Past History）:(?P<既往史>.+)\n' \
            '目前使用的药物（At Present The Drugs）：（含我院用药情况及患者提供的用药情况）(?P<目前使用药物>.+)\n' \
            '成瘾药物\(Drug Addiction\):(?P<成瘾药物>.+)\n' \
            '个人史（Personal History）:(?P<个人史>.+?)\n' \
            '(?P<menstrual>.*)' \
            '婚育史（Obstetrical History）:(?P<婚育史>.+)\n' \
            '家族史（Family History）:(?P<家族史>.+)\n' \
            '体格检查（Physical Examination）：(?P<体格检查>.+)\n' \
            '辅助检查（Diagnostic Examination）：(?P<辅助检查>.+)\n' \
            '营养风险筛查\(Nutritional.+体重指数\(BMI\):(?P<体重指数>.+)\n' \
            '疾病相关评分:\n(?P<疾病相关评分>.+)\n' \
            '营养受损评分:\n(?P<营养受损评分>.+)\n' \
            '年龄评分:(?P<年龄评分>.+)\n' \
            '营养风险评分:(?P<营养风险评分>.+?分).*\n' \
            '是否请营养科会诊:(?P<是否请营养科会诊>.+)\n' \
            '功能评估:\(Function  Accessment\)\n((入院ADL评分:(?P<入院ADL评分>.*))|(入院ADL评分分级:(?P<入院ADL评分分级>.*)))\n' \
            '是否请康复科会诊:(?P<是否请康复科会诊>.+)\n' \
            '心理评估\(Psychological Assessment\)\n护理入院心理评估是否阳性:(?P<护理入院心理评估是否阳性>.*)\n' \
            '是否请心理卫生科会诊:(?P<是否请心理卫生科会诊>.*)\n初步诊断'

    match = re.search(regex, admission_record_text, re.S)
    a = match.groupdict()
    admission_record['主诉'] = a['主诉']
    admission_record['现病史'] = a['现病史']
    admission_record['既往史'] = a['既往史']
    admission_record['目前使用药物'] = a['目前使用药物']
    admission_record['成瘾药物'] = a['成瘾药物']
    admission_record['个人史'] = a['个人史']
    admission_record['婚育史'] = a['婚育史']
    admission_record['家族史'] = a['家族史']
    admission_record['体格检查'] = a['体格检查']
    admission_record['辅助检查'] = a['辅助检查']
    admission_record['营养风险筛查'] = {}
    admission_record['营养风险筛查']['体重指数'] = a['体重指数']
    admission_record['营养风险筛查']['疾病相关评分'] = a['疾病相关评分']
    admission_record['营养风险筛查']['营养受损评分'] = a['营养受损评分']
    admission_record['营养风险筛查']['年龄评分'] = a['年龄评分']
    admission_record['营养风险筛查']['营养风险评分'] = a['营养风险评分']
    admission_record['营养风险筛查']['是否请营养科会诊'] = a['是否请营养科会诊']
    admission_record['功能评估'] = {}
    if '入院ADL评分' in a.keys():
        admission_record['功能评估']['入院ADL评分'] = a['入院ADL评分']
    elif '入院ADL评分分级' in a.keys():
        admission_record['功能评估']['入院ADL评分'] = a['入院ADL评分分级']
    else:
        admission_record['功能评估']['入院ADL评分'] = ''

    admission_record['功能评估']['是否请康复科会诊'] = a['是否请康复科会诊']
    admission_record['心理评估'] = {}
    admission_record['心理评估']['护理入院心理评估是否阳性'] = a['护理入院心理评估是否阳性']
    admission_record['心理评估']['是否请心理卫生科会诊'] = a['是否请心理卫生科会诊']

    regex = '月经史（Menstrual History）:(.+)\n'
    r = re.search(regex, a['menstrual'], re.S)
    admission_record['月经史'] = r.group(1) if r is not None else ''

    regex = '初步诊断\(Diagnosis\).*'
    r = re.search(regex, admission_record_text, re.S)
    diagnosis = r.group()

    regex = '初步诊断\(Diagnosis\)：\n(.+?)医师签名：'
    r = re.search(regex, diagnosis, re.S)
    admission_record['初步诊断'] = r.group(1) if r is not None else ''

    regex = '修正诊断\(Diagnosis\)：\n(.+?)医生签名：'
    r = re.search(regex, diagnosis, re.S)
    admission_record['修正诊断'] = r.group(1) if r is not None else ''

    regex = '补充诊断\(Diagnosis\)：\n(.+?)医生签名：'
    r = re.search(regex, diagnosis, re.S)
    admission_record['补充诊断'] = r.group(1) if r is not None else ''

    # 处理首次病程记录

    regex = '病例特点：(?P<病例特点>.*)\n' \
            '初步诊断：(?P<初步诊断1>.*)\n' \
            '诊断依据：(?P<诊断依据>.*)\n' \
            '鉴别诊断：(?P<鉴别诊断>.*)' \
            '诊疗计划：(?P<诊疗计划>.*)\n.+(医师签名：|记录医生：)'

    match = re.search(regex, progress_note_text, re.S)
    a = match.groupdict()

    progress_note['病例特点'] = a['病例特点']
    progress_note['初步诊断'] = a['初步诊断1']
    progress_note['诊断依据'] = a['诊断依据']
    progress_note['鉴别诊断'] = a['鉴别诊断']
    progress_note['诊疗计划'] = a['诊疗计划']


    if alive:
        # 处理出院记录
        regex = '入院情况:(?P<入院情况>.+)\n' \
                '住院经过(?P<住院经过>.+)\n' \
                '出院情况:(?P<出院情况>.+)\n' \
                '.*出院医嘱(:|：)(?P<出院医嘱>.+)\n' \
                '健康教育:(?P<健康教育>.+)\n' \
                '随访计划:(?P<随访计划>.+)医师签名：'

        match = re.search(regex, discharge_record_text, re.S)
        a = match.groupdict()

        discharge_record['入院情况'] = a['入院情况']
        discharge_record['住院经过'] = a['住院经过']
        discharge_record['出院情况'] = a['出院情况']
        discharge_record['出院医嘱'] = a['出院医嘱']
        discharge_record['健康教育'] = a['健康教育']
        discharge_record['随访计划'] = a['随访计划']
    else:
        # 处理死亡记录
        regex = '入院时间：(?P<入院时间>.+)' \
                '死亡时间：(?P<死亡时间>.+)\n' \
                '(?P<diagnosis>.*)' \
                '入院情况：(?P<入院情况>.+)\n' \
                '诊疗经过：(?P<诊疗经过>.+)\n' \
                '死亡原因：(?P<死亡原因>.+)\n' \
                '死亡诊断：(?P<死亡诊断>.+)医师签名：'

        match = re.search(regex, death_record_text, re.S)
        a = match.groupdict()

        death_record['入院时间'] = a['入院时间']
        death_record['死亡时间'] = a['死亡时间']
        death_record['入院情况'] = a['入院情况']
        death_record['诊疗经过'] = a['诊疗经过']
        death_record['死亡原因'] = a['死亡原因']
        death_record['死亡诊断'] = a['死亡诊断']
        diagnosis = a['diagnosis']
        r = diagnosis.split('入院诊断：')
        if len(r) > 1:
            death_record['入院诊断'] = r[1]
        else:
            death_record['入院诊断'] = ''

    # 合并入院记录、首次病程记录、出院记录、死亡记录到medical_record中

    medical_record = dict()
    medical_record['入院记录'] = admission_record
    medical_record['首次病程记录'] = progress_note
    if alive:
        medical_record['出院记录'] = discharge_record
    else:
        medical_record['死亡记录'] = death_record

    trim_dict_values(medical_record)
    trim_dict_values(medical_record['入院记录'])
    trim_dict_values(medical_record['首次病程记录'])
    if alive:
        trim_dict_values(medical_record['出院记录'])
    else:
        trim_dict_values(medical_record['死亡记录'])

    return medical_record

    pass


def trim_dict_values(dic):
    for key in dic:
        value = dic[key]
        if isinstance(value, str):
            dic[key] = value.strip(' \n')
